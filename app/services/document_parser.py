import os
from pypdf import PdfReader
from docx import Document

class DocumentParser:
    """
    Lớp tiện ích dùng để phân tích và trích xuất nội dung văn bản từ các định dạng file PDF, DOCX, TXT.
    """

    @staticmethod
    def parse_pdf(file_path: str) -> str:
        reader = PdfReader(file_path)
        text = []
        for i, page in enumerate(reader.pages):
            page_text = page.extract_text()
            if page_text:
                text.append(page_text)
        return "\n\n".join(text)

    @staticmethod
    def parse_docx(file_path: str) -> str:
        doc = Document(file_path)
        text = []
        for paragraph in doc.paragraphs:
            if paragraph.text.strip():
                text.append(paragraph.text)
        # Đọc thêm văn bản từ bảng biểu nếu có
        for table in doc.tables:
            for row in table.rows:
                row_text = [cell.text.strip() for cell in row.cells if cell.text.strip()]
                if row_text:
                    text.append(" | ".join(row_text))
        return "\n".join(text)

    @staticmethod
    def parse_txt(file_path: str) -> str:
        # Thử đọc với utf-8, nếu lỗi thử decode với latin-1 hoặc ignore lỗi decode
        try:
            with open(file_path, "r", encoding="utf-8") as f:
                return f.read()
        except UnicodeDecodeError:
            with open(file_path, "r", encoding="latin-1", errors="replace") as f:
                return f.read()

    @classmethod
    def parse(cls, file_path: str) -> str:
        """
        Nhận diện đuôi file và chuyển tiếp đến hàm xử lý tương ứng.
        """
        ext = os.path.splitext(file_path)[1].lower()
        if ext == ".pdf":
            return cls.parse_pdf(file_path)
        elif ext in (".docx", ".doc"):
            return cls.parse_docx(file_path)
        elif ext == ".txt":
            return cls.parse_txt(file_path)
        else:
            raise ValueError(f"Định dạng file {ext} không được hỗ trợ. Chỉ hỗ trợ PDF, DOCX, TXT.")
